import io

import docx
import pytest
from pypdf import PdfWriter

from app.services.text_extraction import (
    EmptyExtractionError,
    UnsupportedFileTypeError,
    extract_text,
    get_file_type,
)


def make_pdf_bytes(text: str) -> bytes:
    """Builds a minimal real PDF in memory. Note: PdfWriter alone can't easily
    embed arbitrary text without a page content stream, so for this test we
    build via reportlab-free approach using pypdf's blank page + we accept
    that extracted text may be empty for a truly blank page - see the
    dedicated 'blank pdf' test below for that behavior."""
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_docx_bytes(paragraphs: list[str]) -> bytes:
    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_get_file_type_accepts_pdf_and_docx():
    assert get_file_type("resume.pdf") == "pdf"
    assert get_file_type("resume.docx") == "docx"


def test_get_file_type_rejects_unsupported_extension():
    with pytest.raises(UnsupportedFileTypeError):
        get_file_type("resume.txt")


def test_get_file_type_rejects_no_extension():
    with pytest.raises(UnsupportedFileTypeError):
        get_file_type("resume")


def test_extract_text_from_docx_pulls_all_paragraphs():
    content = make_docx_bytes([
        "Jane Doe",
        "Software Engineer with 5 years of experience in backend systems.",
        "Skills: Python, FastAPI, PostgreSQL, Docker",
    ])
    file_type, text = extract_text("resume.docx", content)

    assert file_type == "docx"
    assert "Jane Doe" in text
    assert "FastAPI" in text


def test_extract_text_from_docx_includes_table_content():
    document = docx.Document()
    document.add_paragraph("Experience")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Company"
    table.rows[0].cells[1].text = "TechCorp Inc"
    buf = io.BytesIO()
    document.save(buf)

    file_type, text = extract_text("resume.docx", buf.getvalue())
    assert "TechCorp Inc" in text


def test_extract_text_rejects_blank_pdf():
    blank_pdf = make_pdf_bytes("")
    with pytest.raises(EmptyExtractionError):
        extract_text("resume.pdf", blank_pdf)


def test_extract_text_rejects_too_short_docx():
    content = make_docx_bytes(["Hi"])
    with pytest.raises(EmptyExtractionError):
        extract_text("resume.docx", content)
