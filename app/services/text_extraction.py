import io

import docx
from pypdf import PdfReader

from app.utils.exceptions import AppError


class UnsupportedFileTypeError(AppError):
    status_code = 415


class EmptyExtractionError(AppError):
    status_code = 422


SUPPORTED_TYPES = {"pdf", "docx"}


def get_file_type(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_TYPES:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '.{ext}'. Only PDF and DOCX are accepted."
        )
    return ext


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages_text).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables often hold skills/experience in resume templates - don't skip them
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs).strip()


def extract_text(filename: str, file_bytes: bytes) -> tuple[str, str]:
    """Returns (file_type, extracted_text). Raises AppError subclasses on failure."""
    file_type = get_file_type(filename)

    if file_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
    else:
        text = extract_text_from_docx(file_bytes)

    if not text or len(text.strip()) < 20:
        raise EmptyExtractionError(
            "Could not extract meaningful text from this file. "
            "It may be scanned/image-based, corrupted, or empty."
        )

    return file_type, text
